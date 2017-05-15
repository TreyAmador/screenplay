# a pythonic screenplay formatter
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re, os


def open_read(path):
    if path == '':
        path = 'script.docx'
    elif not path.endswith('.docx'):
        path += '.docx'
    if path not in os.listdir():
        return print('That file does not exist.')
    return Document(path), path


def open_write():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(12)
    return doc


def split_bracketed(paragraph,left,right):
    return [x for x in re.split(left+'|'+right,paragraph) if x]


def description(doc,paragraph):
    desc_f = doc.add_paragraph(paragraph)
    desc_f.paragraph_format.left_indent = Inches(0.5)
    desc_f.style = doc.styles['Normal']


def is_subheader(header):
    return header == 'INT' or header == 'EXT' or header == 'SUB'


def heading(doc,header,body):
    head_f = doc.add_paragraph(header+'. '+body.strip().upper())
    head_f.paragraph_format.left_indent = Inches(0.5)
    head_f.style = doc.styles['Normal']


def style_parenthetical():
    pass


def style_paragraph(doc,text,indent,carriage):
    fmt = doc.add_paragraph(text)
    fmt.paragraph_format.left_indent = Inches(indent)
    fmt.style = doc.styles['Normal']
    if carriage is not None:
        fmt.paragraph_format.space_after = Pt(carriage)
    return fmt


def dialogue(doc,header,paragraph):
    fmt_h = style_paragraph(doc,header,2.6,0)
    if paragraph.startswith('('):
        delim = paragraph.strip('(').split(')')
        fmt_a = style_paragraph(doc,'('+delim[0]+')',2.1,0)
        paragraph = delim[1].strip()
    style_paragraph(doc,paragraph,1.5,None)


def convert(path):
    read, path = open_read(path)
    write = open_write()
    for para_obj in read.paragraphs:
        paragraph = split_bracketed(para_obj.text,'<','>')
        if len(paragraph) == 1:
            description(write,paragraph[0])
        elif len(paragraph) >= 2:
            header = paragraph[0].strip().upper()
            if is_subheader(header):
                heading(write,header,paragraph[1])
            elif header == 'TRAN':
                pass
            else:
                dialogue(write,header,paragraph[1].strip())
    directory = path.split('/')
    if len(directory) > 1:
        write.save('/'.join(directory[:-1])+'/format_'+directory[-1])
    else:
        write.save('format_'+path)


def driver():
    path = ''
    while path != 'exit':
        prompt = 'Enter document to convert, "exit" to end, and ' \
                    '"script.docx" is default.\n'
        path = input(prompt)
        path = path.strip().lower()
        if path == 'exit':
            return print('Program ended')
        convert(path)


driver()

# A python script that reads and formats scripts
